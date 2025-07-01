from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Manual Técnico del Proyecto IVACO', 0)

# 1. Descripción General
doc.add_heading('1. Descripción General', level=1)
doc.add_paragraph('IVACO es una aplicación web desarrollada en Django para la gestión y control del estado de declaración del IVA en Colombia. Permite a los usuarios registrar, consultar y administrar facturas, así como visualizar estadísticas y reportes en PDF.')

# 2. Estructura del Proyecto
doc.add_heading('2. Estructura del Proyecto', level=1)
proyecto = [
    '- asistencia/: Aplicación principal de Django, contiene la lógica de negocio, vistas, modelos, formularios y plantillas.',
    '- IVACO/: Configuración global del proyecto Django (settings, urls, wsgi, asgi).',
    '- static/: Archivos estáticos (CSS, imágenes, íconos).',
    '- templates/: Plantillas HTML organizadas por módulos.',
    '- db.sqlite3: Base de datos SQLite por defecto.',
    '- manage.py: Script de administración de Django.'
]
for linea in proyecto:
    doc.add_paragraph(linea, style='List Bullet')

# 3. Instalación y Configuración
doc.add_heading('3. Instalación y Configuración', level=1)
doc.add_paragraph('Requisitos:')
requisitos = [
    'Python 3.10+',
    'pip',
    'Virtualenv (opcional pero recomendado)'
]
for r in requisitos:
    doc.add_paragraph(r, style='List Bullet')
doc.add_paragraph('Instalación:')
instalacion = [
    '1. Clona el repositorio y navega al directorio del proyecto.',
    '2. Crea un entorno virtual:',
    '   python -m venv venv',
    '   venv\\Scripts\\activate',
    '3. Instala las dependencias:',
    '   pip install -r requirements.txt',
    '4. Realiza las migraciones:',
    '   python manage.py migrate',
    '5. Crea un superusuario:',
    '   python manage.py createsuperuser',
    '6. Ejecuta el servidor:',
    '   python manage.py runserver'
]
for i in instalacion:
    doc.add_paragraph(i, style=None)

# 4. Estructura de Archivos Clave
doc.add_heading('4. Estructura de Archivos Clave', level=1)
doc.add_paragraph('Modelos (asistencia/models.py):')
doc.add_paragraph('Define los modelos principales:', style=None)
modelos = [
    '- Usuario: Extiende el modelo de usuario de Django.',
    '- FacturaSubida: Almacena la información de las facturas subidas por los usuarios.'
]
for m in modelos:
    doc.add_paragraph(m, style='List Bullet')
doc.add_paragraph('Vistas (asistencia/views.py):')
vistas = [
    '- Registro y autenticación de usuarios.',
    '- Subida, edición, eliminación y consulta de facturas.',
    '- Generación de reportes PDF.',
    '- Visualización de gráficas y estadísticas.',
    '- Administración de usuarios (CRUD para administradores).'
]
for v in vistas:
    doc.add_paragraph(v, style='List Bullet')
doc.add_paragraph('Formularios (asistencia/forms.py):')
doc.add_paragraph('Define formularios personalizados para autenticación y gestión de usuarios.', style=None)
doc.add_paragraph('Plantillas (asistencia/templates/):')
plantillas = [
    '- plantilla.html: Base para heredar estilos y estructura.',
    '- paginas/: Páginas públicas y de información.',
    '- ivapp/: Vistas principales de usuario (facturas, gráficas, perfil).',
    '- ingresos/: Administración de usuarios.'
]
for p in plantillas:
    doc.add_paragraph(p, style='List Bullet')
doc.add_paragraph('Archivos Estáticos (static/):')
estaticos = [
    '- css/: Hojas de estilo personalizadas.',
    '- images/: Imágenes usadas en la interfaz.',
    '- icons/: Íconos SVG y PNG.'
]
for e in estaticos:
    doc.add_paragraph(e, style='List Bullet')

# 5. Funcionalidades Principales
doc.add_heading('5. Funcionalidades Principales', level=1)
funcs = [
    '- Registro y Login: Usuarios pueden registrarse y autenticarse.',
    '- Gestión de Facturas: Subida, edición, eliminación y consulta de facturas.',
    '- Reportes PDF: Descarga de reportes de facturas en formato PDF.',
    '- Gráficas: Visualización de estadísticas de facturación.',
    '- Modo Oscuro/Claro: Interfaz con soporte para temas.',
    '- Administración: CRUD de usuarios para administradores.'
]
for f in funcs:
    doc.add_paragraph(f, style='List Bullet')

# 6. Seguridad
doc.add_heading('6. Seguridad', level=1)
segs = [
    '- Uso de decoradores @login_required para proteger vistas.',
    '- Validación de datos en formularios y vistas.',
    '- Protección CSRF en formularios.',
    '- Restablecimiento de contraseña vía email.'
]
for s in segs:
    doc.add_paragraph(s, style='List Bullet')

# 7. Personalización y Estilos
doc.add_heading('7. Personalización y Estilos', level=1)
pers = [
    '- Bootstrap 5 para diseño responsivo.',
    '- Hojas de estilo personalizadas en static/css/.',
    '- Soporte para modo oscuro mediante clases CSS y JavaScript.'
]
for p in pers:
    doc.add_paragraph(p, style='List Bullet')

# 8. Despliegue
doc.add_heading('8. Despliegue', level=1)
doc.add_paragraph('Para producción, se recomienda:')
despliegue = [
    '- Configurar un servidor web (Nginx, Apache) y WSGI (Gunicorn, uWSGI).',
    '- Usar una base de datos robusta (PostgreSQL, MySQL).',
    '- Configurar variables de entorno para claves y settings sensibles.',
    '- Servir archivos estáticos y media correctamente.'
]
for d in despliegue:
    doc.add_paragraph(d, style='List Bullet')

# 9. Contacto y Soporte
doc.add_heading('9. Contacto y Soporte', level=1)
doc.add_paragraph('Para dudas técnicas, contactar al equipo de desarrollo o revisar la documentación oficial de Django: https://docs.djangoproject.com/')

doc.save('Manual_Tecnico_IVACO.docx')
